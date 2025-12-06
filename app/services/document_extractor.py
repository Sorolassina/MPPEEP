"""
Service pour extraire le texte des documents uploadés
"""

import io
from pathlib import Path
from typing import Optional
from fastapi import UploadFile

from app.core.logging_config import get_logger

logger = get_logger(__name__)


class DocumentExtractor:
    """Service pour extraire le texte de différents types de documents"""
    
    @staticmethod
    async def extract_text(file: UploadFile) -> Optional[str]:
        """
        Extrait le texte d'un fichier uploadé
        
        Args:
            file: Fichier uploadé
            
        Returns:
            Texte extrait ou None en cas d'erreur
            
        Raises:
            Exception: Si l'extraction échoue (pour que l'endpoint puisse gérer)
        """
        try:
            if not file.filename:
                raise ValueError("Nom de fichier manquant")
            
            filename = file.filename.lower()
            
            # PDF
            if filename.endswith('.pdf'):
                result = await DocumentExtractor._extract_pdf(file)
                if result is None:
                    raise Exception("Impossible d'extraire le texte du PDF")
                return result
            
            # Word documents
            elif filename.endswith(('.doc', '.docx')):
                result = await DocumentExtractor._extract_word(file)
                if result is None:
                    raise Exception("Impossible d'extraire le texte du document Word")
                return result
            
            # Text files
            elif filename.endswith(('.txt', '.md')):
                result = await DocumentExtractor._extract_text_file(file)
                if result is None:
                    raise Exception("Impossible d'extraire le texte du fichier")
                return result
            
            else:
                raise ValueError(f"Type de fichier non supporté: {filename}")
                
        except (ImportError, ValueError, Exception) as e:
            # Re-lancer les erreurs pour que l'endpoint puisse les gérer
            logger.error(f"Erreur lors de l'extraction du texte: {e}", exc_info=True)
            raise
    
    @staticmethod
    async def _extract_pdf(file: UploadFile) -> Optional[str]:
        """Extrait le texte d'un PDF"""
        try:
            # Essayer d'importer PyPDF2 ou pdfplumber
            has_pypdf2 = False
            has_pdfplumber = False
            
            try:
                import PyPDF2
                has_pypdf2 = True
            except ImportError:
                pass
            
            try:
                import pdfplumber
                has_pdfplumber = True
            except ImportError:
                pass
            
            if not has_pypdf2 and not has_pdfplumber:
                error_msg = "Aucune bibliothèque PDF disponible. Installez PyPDF2 (pip install PyPDF2) ou pdfplumber (pip install pdfplumber)."
                logger.error(error_msg)
                raise ImportError(error_msg)
            
            content = await file.read()
            await file.seek(0)
            
            # Essayer avec PyPDF2 d'abord
            if has_pypdf2:
                try:
                    import PyPDF2
                    pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
                    text = ""
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                    if text.strip():
                        return text.strip()
                except Exception as e:
                    logger.warning(f"PyPDF2 a échoué, essai avec pdfplumber: {e}")
            
            # Essayer avec pdfplumber
            if has_pdfplumber:
                try:
                    import pdfplumber
                    text = ""
                    with pdfplumber.open(io.BytesIO(content)) as pdf:
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n"
                    if text.strip():
                        return text.strip()
                except Exception as e:
                    logger.error(f"Erreur avec pdfplumber: {e}")
                    raise Exception(f"Impossible de lire le PDF: {str(e)}. Le fichier est peut-être corrompu ou protégé par mot de passe.")
            
            # Si aucune méthode n'a fonctionné
            raise Exception("Impossible d'extraire le texte du PDF. Le fichier est peut-être corrompu ou protégé.")
                
        except ImportError:
            raise  # Re-lancer les erreurs d'import
        except Exception as e:
            logger.error(f"Erreur lors de l'extraction PDF: {e}", exc_info=True)
            raise  # Re-lancer pour que l'endpoint puisse gérer
    
    @staticmethod
    async def _extract_word(file: UploadFile) -> Optional[str]:
        """Extrait le texte d'un document Word"""
        try:
            try:
                from docx import Document
            except ImportError:
                error_msg = "python-docx n'est pas installé. Installez-le avec: pip install python-docx"
                logger.error(error_msg)
                raise ImportError(error_msg)
            
            content = await file.read()
            await file.seek(0)
            
            try:
                doc = Document(io.BytesIO(content))
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                
                return text.strip()
            except Exception as e:
                logger.error(f"Erreur lors du parsing du document Word: {e}")
                raise Exception(f"Impossible de lire le document Word: {str(e)}. Le fichier est peut-être corrompu ou protégé.")
            
        except ImportError:
            raise  # Re-lancer les erreurs d'import
        except Exception as e:
            logger.error(f"Erreur lors de l'extraction Word: {e}", exc_info=True)
            raise  # Re-lancer pour que l'endpoint puisse gérer
    
    @staticmethod
    async def _extract_text_file(file: UploadFile) -> Optional[str]:
        """Extrait le texte d'un fichier texte"""
        try:
            content = await file.read()
            # Essayer différentes encodages
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    text = content.decode(encoding)
                    return text.strip()
                except UnicodeDecodeError:
                    continue
            
            # Si aucun encodage ne fonctionne, essayer avec errors='ignore'
            return content.decode('utf-8', errors='ignore').strip()
            
        except Exception as e:
            logger.error(f"Erreur lors de l'extraction texte: {e}")
            return None
    
    @staticmethod
    def extract_text_from_content(content: bytes, filename: str) -> str:
        """
        Extrait le texte directement depuis le contenu (bytes) et le nom de fichier
        
        Args:
            content: Contenu du fichier en bytes
            filename: Nom du fichier (pour déterminer le type)
            
        Returns:
            Texte extrait
            
        Raises:
            Exception: Si l'extraction échoue
        """
        try:
            logger.info(f"🔍 [DocumentExtractor] Début extraction pour: {filename} ({len(content)} bytes)")
            
            if not filename:
                raise ValueError("Nom de fichier manquant")
            
            filename_lower = filename.lower()
            logger.info(f"🔍 [DocumentExtractor] Extension détectée: {Path(filename_lower).suffix}")
            
            # PDF
            if filename_lower.endswith('.pdf'):
                logger.info(f"🔍 [DocumentExtractor] Traitement PDF: {filename}")
                return DocumentExtractor._extract_pdf_from_content(content)
            
            # Word documents
            elif filename_lower.endswith(('.doc', '.docx')):
                logger.info(f"🔍 [DocumentExtractor] Traitement Word: {filename}")
                # Vérifier si c'est un fichier .doc (ancien format) ou .docx
                if filename_lower.endswith('.doc') and not filename_lower.endswith('.docx'):
                    # Fichier .doc (ancien format) - python-docx ne peut pas le lire
                    logger.warning(f"⚠️ [DocumentExtractor] Fichier .doc (ancien format) détecté: {filename}")
                    raise ValueError(
                        "Les fichiers .doc (ancien format Microsoft Word) ne sont pas supportés. "
                        "python-docx ne peut lire que les fichiers .docx (format Open XML). "
                        "Veuillez convertir votre fichier .doc en .docx ou utiliser un fichier .docx."
                    )
                logger.info(f"🔍 [DocumentExtractor] Appel _extract_word_from_content pour: {filename}")
                return DocumentExtractor._extract_word_from_content(content)
            
            # Text files
            elif filename_lower.endswith(('.txt', '.md')):
                logger.info(f"🔍 [DocumentExtractor] Traitement fichier texte: {filename}")
                return DocumentExtractor._extract_text_file_from_content(content)
            
            else:
                logger.error(f"❌ [DocumentExtractor] Type de fichier non supporté: {filename_lower}")
                raise ValueError(f"Type de fichier non supporté: {filename_lower}")
                
        except (ImportError, ValueError, Exception) as e:
            logger.error(f"❌ [DocumentExtractor] Erreur lors de l'extraction du texte pour {filename}: {e}", exc_info=True)
            raise
    
    @staticmethod
    def _extract_pdf_from_content(content: bytes) -> str:
        """Extrait le texte d'un PDF depuis le contenu"""
        has_pypdf2 = False
        has_pdfplumber = False
        
        try:
            import PyPDF2
            has_pypdf2 = True
        except ImportError:
            pass
        
        try:
            import pdfplumber
            has_pdfplumber = True
        except ImportError:
            pass
        
        if not has_pypdf2 and not has_pdfplumber:
            error_msg = "Aucune bibliothèque PDF disponible. Installez PyPDF2 (pip install PyPDF2) ou pdfplumber (pip install pdfplumber)."
            logger.error(error_msg)
            raise ImportError(error_msg)
        
        # Essayer avec PyPDF2 d'abord
        if has_pypdf2:
            try:
                import PyPDF2
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                if text.strip():
                    return text.strip()
            except Exception as e:
                logger.warning(f"PyPDF2 a échoué, essai avec pdfplumber: {e}")
        
        # Essayer avec pdfplumber
        if has_pdfplumber:
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(io.BytesIO(content)) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                if text.strip():
                    return text.strip()
            except Exception as e:
                logger.error(f"Erreur avec pdfplumber: {e}")
                raise Exception(f"Impossible de lire le PDF: {str(e)}. Le fichier est peut-être corrompu ou protégé par mot de passe.")
        
        raise Exception("Impossible d'extraire le texte du PDF. Le fichier est peut-être corrompu ou protégé.")
    
    @staticmethod
    def _extract_word_from_content(content: bytes) -> str:
        """Extrait le texte d'un document Word depuis le contenu"""
        try:
            from docx import Document
        except ImportError:
            error_msg = "python-docx n'est pas installé. Installez-le avec: pip install python-docx"
            logger.error(error_msg)
            raise ImportError(error_msg)
        
        try:
            # python-docx ne peut lire que les fichiers .docx (format Open XML)
            # Pour les anciens fichiers .doc, il faudrait utiliser une autre bibliothèque
            doc = Document(io.BytesIO(content))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            result = text.strip()
            if not result:
                raise Exception("Le document Word semble vide ou ne contient pas de texte extractible.")
            
            logger.info(f"✅ Texte Word extrait: {len(result)} caractères")
            return result
        except Exception as e:
            error_msg = str(e)
            logger.error(f"❌ Erreur lors du parsing du document Word: {error_msg}", exc_info=True)
            
            # Vérifier si c'est un fichier .doc (ancien format)
            if "docx.opc.exceptions" in error_msg or "Invalid file format" in error_msg or "not a zip file" in error_msg.lower():
                raise Exception(f"Le fichier semble être au format .doc (ancien format). python-docx ne peut lire que les fichiers .docx (format Open XML). Veuillez convertir le fichier en .docx ou utiliser un fichier .docx.")
            
            raise Exception(f"Impossible de lire le document Word: {error_msg}. Le fichier est peut-être corrompu, protégé par mot de passe, ou au format .doc (non supporté).")
    
    @staticmethod
    def _extract_text_file_from_content(content: bytes) -> str:
        """Extrait le texte d'un fichier texte depuis le contenu"""
        try:
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    text = content.decode(encoding)
                    result = text.strip()
                    if result:
                        logger.info(f"✅ Texte extrait avec encodage {encoding}: {len(result)} caractères")
                        return result
                except UnicodeDecodeError:
                    continue
            
            # Si aucun encodage ne fonctionne, essayer avec errors='ignore'
            result = content.decode('utf-8', errors='ignore').strip()
            if not result:
                raise Exception("Le fichier texte semble vide ou ne contient pas de texte extractible.")
            
            logger.info(f"✅ Texte extrait avec utf-8 (errors='ignore'): {len(result)} caractères")
            return result
        except Exception as e:
            logger.error(f"❌ Erreur lors de l'extraction du fichier texte: {e}", exc_info=True)
            raise Exception(f"Impossible d'extraire le texte du fichier: {str(e)}")

